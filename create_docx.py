from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models import GroupExercise
from typing import Optional
import os

def create_docx_from_group(group: GroupExercise, output_filename: Optional[str] = None):
    """
    Tạo file Word từ một GroupExercise
    
    Args:
        group: GroupExercise object
        output_filename: Tên file output (nếu None sẽ tự động tạo tên)
    """
    # Tạo document mới
    doc = Document()
    
    # Tạo tên file nếu không có
    if not output_filename:
        base_name = group.baseExercise.baseName.replace(" ", "_")
        output_filename = f"exercise_doc/{base_name}_group.docx"
    else:
        # Đảm bảo file được save vào folder exercise_doc
        if not output_filename.startswith("exercise_doc/"):
            output_filename = f"exercise_doc/{output_filename}"
    
    # Thêm tiêu đề chính
    title = doc.add_heading(f'Nhóm bài tập: {group.baseExercise.baseName}', 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thêm bài tập cơ bản
    doc.add_heading('Bài tập cơ bản:', level=2)
    p1 = doc.add_paragraph()
    p1.add_run('Tên: ').bold = True
    p1.add_run(group.baseExercise.name)
    
    p2 = doc.add_paragraph()
    p2.add_run('Tiêu đề: ').bold = True
    p2.add_run(group.baseExercise.title)
    
    p3 = doc.add_paragraph()
    p3.add_run('Nội dung:').bold = True
    doc.add_paragraph(group.baseExercise.content)
    
    # Thêm bài giải nếu có
    if group.solutionExercise:
        doc.add_heading('Bài giải:', level=2)
        p4 = doc.add_paragraph()
        p4.add_run('Tên: ').bold = True
        p4.add_run(group.solutionExercise.name)
        
        p5 = doc.add_paragraph()
        p5.add_run('Tiêu đề: ').bold = True
        p5.add_run(group.solutionExercise.title)
        
        p6 = doc.add_paragraph()
        p6.add_run('Nội dung:').bold = True
        doc.add_paragraph(group.solutionExercise.content)
    else:
        doc.add_heading('Bài giải:', level=2)
        doc.add_paragraph('Không có bài giải')
    
    # Thêm các bài tập mở rộng
    if group.extendedExercises:
        doc.add_heading('Bài tập mở rộng:', level=2)
        for i, exercise in enumerate(group.extendedExercises, 1):
            doc.add_heading(f'{i}. {exercise.name}', level=3)
            p7 = doc.add_paragraph()
            p7.add_run('Tiêu đề: ').bold = True
            p7.add_run(exercise.title)
            
            p8 = doc.add_paragraph()
            p8.add_run('Nội dung:').bold = True
            doc.add_paragraph(exercise.content)
            doc.add_paragraph('')  # Thêm khoảng trống
    else:
        doc.add_heading('Bài tập mở rộng:', level=2)
        doc.add_paragraph('Không có bài tập mở rộng')
    
    # Thêm thông tin tổng quan
    doc.add_heading('Thông tin tổng quan:', level=1)
    doc.add_paragraph(f'• Số bài tập mở rộng: {len(group.extendedExercises)}')
    doc.add_paragraph(f'• Có bài giải: {"Có" if group.solutionExercise else "Không"}')
    doc.add_paragraph(f'• Base name: {group.baseExercise.baseName}')
    
    # Lưu file
    try:
        doc.save(output_filename)
        print(f"Đã tạo file thành công: {output_filename}")
        return True
    except Exception as e:
        print(f"Lỗi khi tạo file: {str(e)}")
        return False

def create_docx_from_groups(groups: list, output_filename: str = "exercise_doc/all_groups.docx"):
    """
    Tạo file Word từ danh sách GroupExercise
    
    Args:
        groups: List các GroupExercise
        output_filename: Tên file output
    """
    # Đảm bảo file được save vào folder exercise_doc
    if not output_filename.startswith("exercise_doc/"):
        output_filename = f"exercise_doc/{output_filename}"
    
    doc = Document()
    
    # Thêm tiêu đề chính
    title = doc.add_heading('Tất cả nhóm bài tập', 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thêm từng nhóm
    for i, group in enumerate(groups, 1):
        doc.add_heading(f'Nhóm {i}: {group.baseExercise.baseName}', level=2)
        
        # Thêm thông tin tóm tắt
        doc.add_paragraph(f'• Base exercise: {group.baseExercise.name}')
        doc.add_paragraph(f'• Solution: {"Có" if group.solutionExercise else "Không"}')
        doc.add_paragraph(f'• Extended exercises: {len(group.extendedExercises)}')
        
        # Thêm nội dung base exercise
        p = doc.add_paragraph()
        p.add_run('Nội dung bài cơ bản:').bold = True
        doc.add_paragraph(group.baseExercise.content[:200] + "..." if len(group.baseExercise.content) > 200 else group.baseExercise.content)
        
        doc.add_paragraph('')  # Thêm khoảng trống
    
    # Lưu file
    try:
        doc.save(output_filename)
        print(f"Đã tạo file thành công: {output_filename}")
        return True
    except Exception as e:
        print(f"Lỗi khi tạo file: {str(e)}")
        return False

# Test function
if __name__ == "__main__":
    # Import để test
    import sys
    sys.path.append('..')
    from models import GroupExercise, ObjExercise
    
    # Tạo test data
    base_exercise = ObjExercise(name="Bài 1", title="Bài tập cơ bản", content="Nội dung bài tập cơ bản")
    solution_exercise = ObjExercise(name="Bài 1 - Giải", title="Giải bài tập", content="Lời giải chi tiết")
    extended1 = ObjExercise(name="Bài 1.1", title="Bài tập mở rộng 1", content="Nội dung mở rộng 1")
    extended2 = ObjExercise(name="Bài 1.2", title="Bài tập mở rộng 2", content="Nội dung mở rộng 2")
    
    group = GroupExercise(
        baseExercise=base_exercise,
        solutionExercise=solution_exercise,
        extendedExercises=[extended1, extended2]
    )
    
    # Test tạo file
    create_docx_from_group(group, "test_group.docx") 