from typing import List
from models.exercise.exercise import Exercise
from models.exercise.extended_exercise import ExtendedExercise
from docx import Document
from write_doc.utils_write_doc import create_table, write_list_section, write_description_section, add_style_paragraph, fix_column_widths, add_style_text, add_style_cell
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT


def write_extend_exercise(doc: Document, baseExercise: List[List[ExtendedExercise]]):# type: ignore
    header = doc.add_heading(f"Bài tập mở rộng", level=2)
    add_style_paragraph(header, {
            'rgb_color': RGBColor(255, 0, 0),
            'alignment': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'bold': True,
            'font_size': 20
        })
    doc.add_paragraph()
    for i, exercise in enumerate(baseExercise, 1):
        write_exercise(doc, exercise, str(i))

def write_exercise(doc: Document, baseExercise: List[ExtendedExercise], prefix: str):# type: ignore
    for i, exercise in enumerate(baseExercise, 1):
        header = doc.add_heading(f"Bài {prefix}.{i}: {exercise.title}", level=3)
        
        add_style_paragraph(header, {
            'rgb_color': RGBColor(0, 111, 192),
            'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'bold': True,
            'space_before':20
        })

        write_description_section(doc, exercise.description_section, {
            'font_size': 14,
            'alignment': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        })
        style_title = {
                'bold': True
            }
        style_list = {
            'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'left_indent': 0.82
        }
        if(len(exercise.input_section) > 0):
            input_para = doc.add_paragraph("Input:")
            add_style_paragraph(input_para, style_title)
        write_list_section(doc, exercise.input_section, 'List Bullet', style_list)
        
        # Output section with bold label
        if(len(exercise.output_section) > 0):
            output_para = doc.add_paragraph("Output:")
          
            add_style_paragraph(output_para, style_title)
        write_list_section(doc, exercise.output_section, 'List Bullet', style_list)

        example_para = doc.add_paragraph("Example:")
        add_style_paragraph(example_para, style_title)

        table = create_table(doc, 2, 2)
        fix_column_widths(table, [2, 2])
        table_style = {
        'alignment': WD_PARAGRAPH_ALIGNMENT.CENTER,
        'bold': True,
        'vertical_alignment': WD_CELL_VERTICAL_ALIGNMENT.CENTER,
         }
        add_style_cell(table.cell(0, 0), 'Input', table_style)
        add_style_cell(table.cell(0, 1), 'Output', table_style)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        table_style = {
            'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'vertical_alignment': WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            'font_size': 14
        }
        input_sample = exercise.input_sample
        for j in range(len(input_sample)):
            if(j == 0):
                add_style_cell(table.cell(1, 0), input_sample[j], table_style)
            else:
                add_style_text(table.cell(1, 0), input_sample[j], table_style)
        
        table_style = {
            'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'vertical_alignment': WD_CELL_VERTICAL_ALIGNMENT.TOP,
            'font_size': 14
        }
        output_sample = exercise.output_sample
        for j, value in enumerate(output_sample):
            if(j == 0):
                add_style_cell(table.cell(1, 1), str(value), table_style)
            else:
                add_style_text(table.cell(1, 1), str(value), table_style)


