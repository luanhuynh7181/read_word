from typing import List
from models.exercise.solution import Solution
from docx import Document
from write_doc.utils_write_doc import create_table, add_style_paragraph, write_list_section_number, add_style_cell, add_style_text

from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from pygments import lex
from pygments.lexers import PythonLexer
from pygments.token import Token

def write_solution_exercise(doc: Document, solutionExercise: List[Solution]): # type: ignore
    header = doc.add_heading(f"Hướng dẫn giải", level=2)
    add_style_paragraph(header, {
            'rgb_color': RGBColor(255, 0, 0),
            'alignment': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'bold': True,
            'font_size': 20,
            'line_spacing': 25
        })

    for i, exercise in enumerate(solutionExercise, 0):
        header = doc.add_heading(f"Bài {i+1}: {exercise.title}", level=3)
        add_style_paragraph(header, {
            'rgb_color': RGBColor(0, 111, 192),
            'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'bold': True,
            'space_before':20
        })
        
        phantich = doc.add_paragraph("1. Phân tích:")
        add_style_paragraph(phantich, {
            'rgb_color': RGBColor(0, 111, 192),
            'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'bold': True,
            'space_before':15
        })

        write_list_section_number(doc, exercise.solution_textes, {
            'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'left_indent': 0.82
        })
        code = doc.add_paragraph("2. Code tham khảo:",  style="Normal")
        add_style_paragraph(code, {
            'rgb_color': RGBColor(0, 111, 192),
            'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'bold': True
        })
        create_code_table(doc, exercise.code_sample) # type: ignore
    pass

def create_code_table(doc: Document, code: List[str]): # type: ignore
    table = create_table(doc, 1, 1)
    write_code_table(code, table.cell(0, 0))

def write_code_table(code_lines: List[str], cell):
    paragraph = cell.paragraphs[0]
    code = "\n".join(code_lines)

    color_map = {
        Token.Keyword: RGBColor(0, 0, 255),               # Blue - for, def, import
        Token.Name.Function: RGBColor(0, 128, 0),         # Dark Green
        Token.Name.Class: RGBColor(0, 102, 102),          # Teal
        Token.Name.Builtin: RGBColor(102, 0, 153),        # Purple - print, len
        Token.Literal.String: RGBColor(196, 26, 22),      # Red - strings
        Token.Literal.Number: RGBColor(255, 85, 0),       # Orange
        Token.Comment: RGBColor(0, 128, 0),               # Green - comments
        Token.Operator: RGBColor(0, 102, 204),            # Blue-ish
        Token.Punctuation: RGBColor(128, 128, 128),       # Gray
        Token.Name: RGBColor(0, 102, 102),                # Teal (variables)
        Token.Text: RGBColor(80, 80, 80),                 # Dark Gray (space/newline)
    }

    tokens = list(lex(code, PythonLexer()))

    if tokens and tokens[-1][1] == '\n':
        tokens.pop()

    for token_type, token_value in tokens:
        run = paragraph.add_run(token_value)
        rgb = color_map.get(token_type, RGBColor(150, 75, 0))  # fallback màu đẹp
        run.font.color.rgb = rgb
        run.font.name = "Consolas"
