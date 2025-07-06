from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from typing import List
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Cm
import re
def create_table(doc: Document, rows: int, cols: int): # type: ignore
    table = doc.add_table(rows=rows, cols=cols)
    set_table_border(table)
    return table

def set_table_border(table):
    tbl = table._tbl  # XML của bảng
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.get_or_add_tblPr()

    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')         # độ dày viền
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000') # màu đen
        tblBorders.append(border)

    tblPr.append(tblBorders)
 
import re
from docx import Document

def add_paragraph_contain_superscript(doc: Document, item: str, style="Normal", prefix: str = ""):  # type: ignore
    pattern = re.compile(r'10(\d{1,2})(?!\d)')  # match 10 + 1–2 chữ số, không kèm thêm số
    para = doc.add_paragraph(style=style)
    para.add_run(prefix)  # thêm tiền tố nếu có
    last_index = 0
    for match in pattern.finditer(item):
        start, end = match.span()
        full_match = match.group(0)  # ví dụ: "104"
        exponent = match.group(1)    # ví dụ: "4"
        if full_match in ["100", "1000"]:
            continue
        para.add_run(item[last_index:start])
        para.add_run("10")
        for ch in exponent:
            r = para.add_run(ch)
            r.font.superscript = True

        last_index = end
    para.add_run(item[last_index:])
    return para


def write_list_section(doc: Document, listSection: List[str], type: str = 'List Bullet', style =  None):# type: ignore
    for item in listSection:
        para = add_paragraph_contain_superscript(doc, item, type)
        if style is not None:
            add_style_paragraph(para, style)

def write_list_section_number(doc: Document, listSection: List[str], style=None):  # type: ignore
    for i, item in enumerate(listSection, start=1):
        para = add_paragraph_contain_superscript(doc, item, "Normal", f"{i}. ")
        if style is not None:
            add_style_paragraph(para, style)

def write_description_section(doc: Document, descriptionSection: List[str], style =  None):# type: ignore
    for item in descriptionSection:
        para = add_paragraph_contain_superscript(doc, item, "Normal", "    ")
        if style is not None:
            add_style_paragraph(para, style)
        

def add_style(run, obj):

    if 'bold' in obj:
        run.bold = obj['bold']
    
    if 'font_size' not in obj:
        obj['font_size'] = 14
    if 'font_size' in obj:
        run.font.size = Pt(obj['font_size'])

    if "font_name" not in obj:
        obj["font_name"] = "Times New Roman"

    if 'font_name' in obj:
        run.font.name = obj['font_name']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), obj['font_name'])

    if 'rgb_color' in obj:
        run.font.color.rgb = obj['rgb_color']


def add_style_paragraph(para, style):
    if "line_spacing" in style:
        para.paragraph_format.line_spacing = Pt(style["line_spacing"])
    else:
        para.paragraph_format.line_spacing = Pt(17)

    if "space_before" in style:
        para.paragraph_format.space_before = Pt(style["space_before"])
    else:
        para.paragraph_format.space_before = Pt(8)

    if "space_after" in style:
        para.paragraph_format.space_after = Pt(style["space_after"])
    else:
        para.paragraph_format.space_after = Pt(8)

    if 'vertical_alignment' in style:
        para.paragraph_format.vertical_alignment = style['vertical_alignment']

    if "alignment" not in style:
        style["alignment"] = WD_PARAGRAPH_ALIGNMENT.LEFT

    if 'alignment' in style:
        para.alignment = style['alignment']

    if 'left_indent' in style:
        para.paragraph_format.left_indent = Cm(style['left_indent'])

    if 'right_indent' in style:
        para.paragraph_format.right_indent = Cm(style['right_indent'])

    if 'first_line_indent' in style:
        para.paragraph_format.first_line_indent = Pt(style['first_line_indent'])
    for run in para.runs:
        add_style(run, style)

    
def add_style_text(parent, text, style):
    para = parent.add_paragraph(text)
    add_style_paragraph(para, style)

def add_style_cell(cell, text, style):
    cell.text = text
    add_style_paragraph(cell.paragraphs[0], style)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Tạo hoặc lấy <w:tcMar>
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)

    # Tạo hoặc cập nhật <w:left>
    left = OxmlElement('w:left')
    left.set(qn('w:w'), str(200))  # Ví dụ: 283 = 0.5 cm
    left.set(qn('w:type'), 'dxa')
    tcMar.append(left)

    right = OxmlElement('w:right')
    right.set(qn('w:w'), str(200))  # Ví dụ: 283 = 0.5 cm
    right.set(qn('w:type'), 'dxa')
    tcMar.append(right)

def fix_column_widths(table, widths_in_inches):
    # Tắt tự động giãn
    table.autofit = True
    for row in table.rows:
        for i, width in enumerate(widths_in_inches):
            cell = row.cells[i]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'dxa')
            tcW.set(qn('w:w'), str(int(width * 1440)))  # 1 inch = 1440 dxa
            tcPr.append(tcW)